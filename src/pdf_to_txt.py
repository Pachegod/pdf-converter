import PyPDF2
import os
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def encontrar_tesseract():
    """Procura o Tesseract em diferentes locais comuns"""
    possiveis_caminhos = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        os.path.expanduser('~\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe'),
        os.path.expanduser('~\\AppData\\Local\\Programs\\Python\\Python3*\\Scripts\\tesseract.exe'),
        os.path.expanduser('~\\AppData\\Local\\Packages\\PythonSoftwareFoundation.Python.3*\\LocalCache\\local-packages\\Python3*\\Scripts\\tesseract.exe')
    ]
    
    for caminho in possiveis_caminhos:
        if os.path.exists(caminho):
            return caminho
    
    return None

def verificar_idioma_tesseract():
    """Verifica se o arquivo de idioma português está disponível"""
    possiveis_caminhos = [
        r'C:\Program Files\Tesseract-OCR\tessdata\por.traineddata',
        r'C:\Program Files (x86)\Tesseract-OCR\tessdata\por.traineddata',
        os.path.expanduser('~\\AppData\\Local\\Programs\\Tesseract-OCR\\tessdata\\por.traineddata')
    ]
    
    for caminho in possiveis_caminhos:
        if os.path.exists(caminho):
            return True
    
    return False

# Configura o caminho do Tesseract
caminho_tesseract = encontrar_tesseract()
if caminho_tesseract:
    pytesseract.pytesseract.tesseract_cmd = caminho_tesseract
    if not verificar_idioma_tesseract():
        print('AVISO: Arquivo de idioma português não encontrado.')
        print('Por favor:')
        print('1. Baixe o arquivo por.traineddata de: https://github.com/tesseract-ocr/tessdata')
        print('2. Copie o arquivo para a pasta tessdata do Tesseract')
        print('3. O caminho deve ser: C:\\Program Files\\Tesseract-OCR\\tessdata\\por.traineddata')
else:
    print('AVISO: Tesseract não encontrado. O OCR não estará disponível.')
    print('Por favor, instale o Tesseract seguindo as instruções:')
    print('1. Baixe o Tesseract de: https://github.com/UB-Mannheim/tesseract/wiki')
    print('2. Execute o instalador')
    print('3. Marque a opção "Add to system PATH" durante a instalação')
    print('4. Reinicie o terminal após a instalação')

# Configura o caminho do Poppler
POPPLER_PATH = r'C:\Program Files\poppler\Library\bin'

def converter_pdf_para_txt(caminho_pdf, caminho_saida=None, usar_ocr=False):
    """
    Converte um arquivo PDF para TXT
    :param caminho_pdf: Caminho do arquivo PDF
    :param caminho_saida: Caminho onde o arquivo TXT será salvo (opcional)
    :param usar_ocr: Se True, usa OCR para extrair texto de imagens
    :return: Caminho do arquivo TXT gerado
    """
    try:
        # Verifica se o arquivo existe
        if not os.path.exists(caminho_pdf):
            print(f'Erro: O arquivo "{caminho_pdf}" não foi encontrado.')
            print('Verifique se:')
            print('1. O caminho está correto')
            print('2. O arquivo existe no local especificado')
            print('3. Você incluiu a extensão .pdf no nome do arquivo')
            return None

        # Verifica se é um arquivo PDF
        if not caminho_pdf.lower().endswith('.pdf'):
            print('Erro: O arquivo deve ter a extensão .pdf')
            return None

        # Se não foi especificado um caminho de saída, usa o mesmo nome do PDF
        if caminho_saida is None:
            caminho_saida = os.path.splitext(caminho_pdf)[0] + '.txt'

        texto_completo = ''

        if usar_ocr:
            if not caminho_tesseract:
                print('Erro: OCR não disponível. Tesseract não encontrado.')
                return None
            
            if not verificar_idioma_tesseract():
                print('Erro: Arquivo de idioma português não encontrado.')
                print('Por favor, instale o arquivo por.traineddata seguindo as instruções acima.')
                return None
                
            print('Usando OCR para extrair texto das imagens...')
            # Verifica se o Poppler está instalado
            if not os.path.exists(POPPLER_PATH):
                print(f'Erro: Poppler não encontrado em {POPPLER_PATH}')
                print('Por favor, instale o Poppler seguindo as instruções:')
                print('1. Baixe o Poppler de: https://github.com/oschwartz10612/poppler-windows/releases/')
                print('2. Extraia o arquivo ZIP')
                print('3. Copie a pasta para C:\\Program Files\\poppler')
                print('4. Adicione C:\\Program Files\\poppler\\Library\\bin ao PATH do sistema')
                return None

            # Converte PDF para imagens
            imagens = convert_from_path(caminho_pdf, poppler_path=POPPLER_PATH)
            
            # Processa cada página
            for i, imagem in enumerate(imagens):
                print(f'Processando página {i+1} de {len(imagens)}...')
                # Extrai texto da imagem usando OCR
                texto = pytesseract.image_to_string(imagem, lang='por')
                texto_completo += texto + '\n'
        else:
            # Tenta primeiro extrair texto normalmente
            with open(caminho_pdf, 'rb') as arquivo_pdf:
                leitor_pdf = PyPDF2.PdfReader(arquivo_pdf)
                
                # Extrai o texto de todas as páginas
                for pagina in leitor_pdf.pages:
                    texto_completo += pagina.extract_text() + '\n'
            
            # Se não encontrou texto, tenta usar OCR
            if not texto_completo.strip():
                print('Nenhum texto encontrado. Tentando usar OCR...')
                return converter_pdf_para_txt(caminho_pdf, caminho_saida, usar_ocr=True)
        
        # Salva o texto em um arquivo TXT
        with open(caminho_saida, 'w', encoding='utf-8') as arquivo_txt:
            arquivo_txt.write(texto_completo)
        
        print(f'Arquivo TXT criado com sucesso: {caminho_saida}')
        return caminho_saida
            
    except Exception as e:
        print(f'Erro ao converter o PDF: {str(e)}')
        return None

def converter_pdf_para_docx(caminho_pdf, caminho_saida=None, usar_ocr=False):
    """
    Converte um arquivo PDF para DOCX
    :param caminho_pdf: Caminho do arquivo PDF
    :param caminho_saida: Caminho onde o arquivo DOCX será salvo (opcional)
    :param usar_ocr: Se True, usa OCR para extrair texto de imagens
    :return: Caminho do arquivo DOCX gerado
    """
    try:
        # Verifica se o arquivo existe
        if not os.path.exists(caminho_pdf):
            print(f'Erro: O arquivo "{caminho_pdf}" não foi encontrado.')
            print('Verifique se:')
            print('1. O caminho está correto')
            print('2. O arquivo existe no local especificado')
            print('3. Você incluiu a extensão .pdf no nome do arquivo')
            return None

        # Verifica se é um arquivo PDF
        if not caminho_pdf.lower().endswith('.pdf'):
            print('Erro: O arquivo deve ter a extensão .pdf')
            return None

        # Se não foi especificado um caminho de saída, usa o mesmo nome do PDF
        if caminho_saida is None:
            caminho_saida = os.path.splitext(caminho_pdf)[0] + '.docx'

        # Cria um novo documento Word
        doc = Document()
        
        # Configura o estilo padrão
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        if usar_ocr:
            if not caminho_tesseract:
                print('Erro: OCR não disponível. Tesseract não encontrado.')
                return None
            
            if not verificar_idioma_tesseract():
                print('Erro: Arquivo de idioma português não encontrado.')
                print('Por favor, instale o arquivo por.traineddata seguindo as instruções acima.')
                return None
                
            print('Usando OCR para extrair texto das imagens...')
            # Verifica se o Poppler está instalado
            if not os.path.exists(POPPLER_PATH):
                print(f'Erro: Poppler não encontrado em {POPPLER_PATH}')
                print('Por favor, instale o Poppler seguindo as instruções:')
                print('1. Baixe o Poppler de: https://github.com/oschwartz10612/poppler-windows/releases/')
                print('2. Extraia o arquivo ZIP')
                print('3. Copie a pasta para C:\\Program Files\\poppler')
                print('4. Adicione C:\\Program Files\\poppler\\Library\\bin ao PATH do sistema')
                return None

            # Converte PDF para imagens
            imagens = convert_from_path(caminho_pdf, poppler_path=POPPLER_PATH)
            
            # Processa cada página
            for i, imagem in enumerate(imagens):
                print(f'Processando página {i+1} de {len(imagens)}...')
                # Extrai texto da imagem usando OCR
                texto = pytesseract.image_to_string(imagem, lang='por')
                # Adiciona o texto ao documento Word
                paragrafo = doc.add_paragraph(texto)
                # Adiciona uma quebra de página após cada página
                if i < len(imagens) - 1:
                    doc.add_page_break()
        else:
            # Tenta primeiro extrair texto normalmente
            with open(caminho_pdf, 'rb') as arquivo_pdf:
                leitor_pdf = PyPDF2.PdfReader(arquivo_pdf)
                
                # Extrai o texto de todas as páginas
                for i, pagina in enumerate(leitor_pdf.pages):
                    print(f'Processando página {i+1} de {len(leitor_pdf.pages)}...')
                    texto = pagina.extract_text()
                    # Adiciona o texto ao documento Word
                    paragrafo = doc.add_paragraph(texto)
                    # Adiciona uma quebra de página após cada página
                    if i < len(leitor_pdf.pages) - 1:
                        doc.add_page_break()
            
            # Se não encontrou texto, tenta usar OCR
            if not doc.paragraphs:
                print('Nenhum texto encontrado. Tentando usar OCR...')
                return converter_pdf_para_docx(caminho_pdf, caminho_saida, usar_ocr=True)
        
        # Salva o documento
        doc.save(caminho_saida)
        print(f'Arquivo DOCX criado com sucesso: {caminho_saida}')
        return caminho_saida
            
    except Exception as e:
        print(f'Erro ao converter o PDF: {str(e)}')
        return None

if __name__ == '__main__':
    print('=== Conversor de PDF ===')
    print('1. Converter para TXT')
    print('2. Converter para DOCX')
    opcao = input('Escolha a opção (1 ou 2): ').strip()
    
    print('\nExemplo de caminho: C:\\Users\\SeuUsuario\\Downloads\\documento.pdf')
    print('\nDicas:')
    print('1. Use \\\\ ou / para separar pastas')
    print('2. Inclua a extensão .pdf no nome do arquivo')
    print('3. Use o caminho completo do arquivo\n')
    
    caminho_pdf = input('Digite o caminho do arquivo PDF: ').strip()
    caminho_saida = input('Digite o caminho onde deseja salvar o arquivo (pressione Enter para usar o nome padrão): ').strip()
    usar_ocr = input('Deseja usar OCR para extrair texto de imagens? (s/n): ').strip().lower() == 's'
    
    if not caminho_saida:
        caminho_saida = None
    
    if opcao == '1':
        converter_pdf_para_txt(caminho_pdf, caminho_saida, usar_ocr)
    elif opcao == '2':
        converter_pdf_para_docx(caminho_pdf, caminho_saida, usar_ocr)
    else:
        print('Opção inválida!') 